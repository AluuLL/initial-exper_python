from docx import Document
from docxpl import DocxTemplate,InlineImage
import random
import re
#coding=utf-8
d = Document(r'/disk/lulu/autoreport/template.docx')
l = [ paragraph.text.encode('UTF-8') for paragraph in d.paragraphs]

for key,para in enumerate(l):

    if key ==4:
        print str(para)

p = re.compile("\d")
for num,table in enumerate(d.tables):
#    if num == 0:
#   find_target='58'
#   substitute='chenran'
#    for row in table.rows:
#        for cell in row.cells:
#            cell.text = cell.text.replace(find_target,substitute)
       for row in table.rows:
            for cell in row.cells:
                if re.search(p,cell.text):
                    cell.text = "lulu"

d.add_paragraph('last paragraph',style=None)
t1 = d.add_table(rows=2,cols=3,style=None)
print len(t1.rows)
print len(t1.columns)

for row in t1.rows:
    row.cells[0].text = '1'
    print row.cells[0].text


for col in t1.columns:
    col.cells[0].text = '2'
    print col.cells[0].text

d.save('/disk/lulu/autoreport/new.docx')
